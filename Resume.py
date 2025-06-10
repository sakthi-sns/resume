import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import os

# --- Secure Gemini API Key Setup ---
api_key = os.getenv("GEMINI_API_KEY") or "AIzaSyBTPDAVY2fKwg5zvipyIDkzqwybsoMO5qA"
genai.configure(api_key=api_key)

# Load Gemini model
model = genai.GenerativeModel("gemini-2.0-flash")

# --- Helper to generate resume content using Gemini ---
def generate_resume(prompt):
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error generating resume: {e}"

# --- Create DOCX Resume File ---
def create_docx(content):
    doc = Document()
    doc.add_heading("Resume", level=1)
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- Streamlit UI ---
st.set_page_config(page_title="Resume Generator", layout="centered")
st.title("📄 AI Resume Generator")
st.write("Fill in your details to get a professionally formatted resume.")

# --- Input Form ---
with st.form("resume_form"):
    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    location = st.text_input("Location")
    summary = st.text_area("Professional Summary", height=100)
    skills = st.text_area("Skills (comma-separated)")
    experience = st.text_area("Work Experience", height=150)
    education = st.text_area("Education")
    job_role = st.selectbox("Target Job Role", ["Software Engineer", "Data Analyst", "Product Manager", "UX Designer", "Custom"])
    custom_role = ""
    if job_role == "Custom":
        custom_role = st.text_input("Enter Custom Job Role")
    submit = st.form_submit_button("Generate Resume")

# --- Generate Resume on Submit ---
if submit:
    final_role = custom_role if job_role == "Custom" else job_role
    with st.spinner("Generating your resume..."):
        prompt = f"""
        Create a professionally formatted resume for the following details:
        Name: {name}
        Email: {email}
        Phone: {phone}
        Location: {location}
        Professional Summary: {summary}
        Skills: {skills}
        Work Experience: {experience}
        Education: {education}
        Tailor the resume for the role of a {final_role}.
        Make it concise, formal, and attractive for recruiters.
        """
        result = generate_resume(prompt)

    if "Error" in result:
        st.error(result)
    else:
        st.success("Resume generated successfully!")
        st.text_area("Generated Resume Preview", result, height=300)
        docx_file = create_docx(result)
        st.download_button(
            label="📥 Download as DOCX",
            data=docx_file,
            file_name=f"{name.replace(' ', '_')}_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
